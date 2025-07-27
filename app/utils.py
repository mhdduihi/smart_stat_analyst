import pandas as pd
import numpy as np
from scipy import stats
import statsmodels.api as sm
import matplotlib.pyplot as plt
from docx import Document
import pyreadstat
from pingouin import cronbach_alpha
import os

def read_file(file_path):
    """قراءة ملف CSV أو Excel."""
    if file_path.endswith('.csv'):
        return pd.read_csv(file_path)
    elif file_path.endswith('.xlsx'):
        return pd.read_excel(file_path)
    else:
        raise ValueError("نوع الملف غير مدعوم")

def perform_analysis(file_path, questions, hypotheses):
    """إجراء التحليل الإحصائي وتوليد المخرجات."""
    # قراءة البيانات
    df = read_file(file_path)

    # تحديد نوع المتغيرات
    categorical_vars = df.select_dtypes(include=['object']).columns
    numerical_vars = df.select_dtypes(include=['number']).columns

    # الإحصاءات الوصفية
    desc_stats = {}
    for col in categorical_vars:
        desc_stats[col] = df[col].value_counts(normalize=True)
    for col in numerical_vars:
        desc_stats[col] = df[col].describe()

    # تحليل الثبات (Cronbach’s Alpha)
    alpha = None
    if len(numerical_vars) > 1:
        alpha = cronbach_alpha(df[numerical_vars])[0]

    # اختبار الفرضيات (مصفوفة الارتباط)
    corr_matrix = None
    if len(numerical_vars) > 1:
        corr_matrix = df[numerical_vars].corr(method='pearson')

    # ✅ إنشاء الرسوم البيانية بشكل آمن
    os.makedirs('output/charts', exist_ok=True)
    for col in numerical_vars:
        try:
            plt.figure()
            df[col].dropna().hist(color='#1f77b4')
            plt.title(f'توزيع {col}')
            plt.xlabel(col)
            plt.ylabel('التكرار')
            plt.savefig(f'output/charts/{col}_hist.png')
            plt.close()
        except Exception as e:
            print(f"تعذر رسم العمود {col}: {e}")

    # إنشاء تقرير Word
    doc = Document()
    doc.add_heading('تقرير البحث', 0)
    doc.add_paragraph('أسئلة البحث:')
    doc.add_paragraph(questions)
    doc.add_paragraph('الفرضيات:')
    doc.add_paragraph(hypotheses)
    doc.add_heading('الإحصاءات الوصفية', level=1)
    for col, stats in desc_stats.items():
        doc.add_heading(f'المتغير: {col}', level=2)
        doc.add_paragraph(str(stats))
    if alpha is not None:
        doc.add_heading('تحليل الثبات', level=1)
        doc.add_paragraph(f'معامل كرونباخ ألفا: {alpha:.2f}')
    if corr_matrix is not None:
        doc.add_heading('اختبار الفرضيات', level=1)
        doc.add_paragraph('مصفوفة الارتباط (Pearson):')
        doc.add_paragraph(str(corr_matrix))
    doc.save('output/word_report.docx')

    # إنشاء ملخص Excel
    with pd.ExcelWriter('output/excel_summary.xlsx') as writer:
        for col, stats in desc_stats.items():
            stats.to_excel(writer, sheet_name=col[:31])

    # ✅ تصحيح أسماء الأعمدة قبل حفظ SPSS
    df.columns = [col.strip().replace(" ", "_")[:32] for col in df.columns]
    pyreadstat.write_sav(df, 'output/data_export.sav')
